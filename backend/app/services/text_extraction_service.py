"""
Text Extraction Service

This service handles text extraction from various file formats including
PDF, DOCX, TXT, and Markdown files for processing by LightRAG.

Features:
- Support for multiple file formats
- Robust encoding detection for text files
- Clean text preprocessing
- Error handling and fallback methods
"""

import logging
import io
import re
from typing import Union
from pathlib import Path

# Document processing libraries
import docx
import PyPDF2
import pdfplumber

logger = logging.getLogger(__name__)

def extract_text_from_pdf(file_content: bytes) -> str:
    """
    Extract text content from PDF file bytes.
    
    Args:
        file_content: PDF file content as bytes
        
    Returns:
        Extracted text content
        
    Raises:
        Exception: If PDF extraction fails completely
    """
    try:
        text = ""
        
        # Try with pdfplumber first (better for complex layouts)
        try:
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
            if text.strip():
                logger.info(f"Successfully extracted text from PDF using pdfplumber ({len(text)} chars)")
                return text.strip()
        except Exception as e:
            logger.warning(f"pdfplumber extraction failed: {e}")
        
        # Fallback to PyPDF2
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        
        if not text.strip():
            raise Exception("No text could be extracted from the PDF")
        
        logger.info(f"Successfully extracted text from PDF using PyPDF2 ({len(text)} chars)")
        return text.strip()
        
    except Exception as e:
        logger.error(f"Failed to extract text from PDF: {e}")
        raise Exception(f"PDF text extraction failed: {e}")

def extract_text_from_docx(file_content: bytes) -> str:
    """
    Extract text content from Word document bytes.
    
    Args:
        file_content: DOCX file content as bytes
        
    Returns:
        Extracted text content
        
    Raises:
        Exception: If DOCX extraction fails
    """
    try:
        doc = docx.Document(io.BytesIO(file_content))
        text = ""
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        if not text.strip():
            raise Exception("No text could be extracted from the Word document")
        
        logger.info(f"Successfully extracted text from DOCX ({len(text)} chars)")
        return text.strip()
        
    except Exception as e:
        logger.error(f"Failed to extract text from DOCX: {e}")
        raise Exception(f"DOCX text extraction failed: {e}")

def extract_text_from_txt(file_content: bytes) -> str:
    """
    Extract text content from plain text file bytes.
     
    Args:
        file_content: Text file content as bytes
        
    Returns:
        Extracted text content
        
    Raises:
        Exception: If text extraction fails
    """
    try:
        # Try different encodings
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                text = file_content.decode(encoding)
                    
                if text.strip():
                    logger.info(f"Successfully extracted text from TXT using {encoding} ({len(text)} chars)")
                    return text.strip()
            except UnicodeDecodeError:
                continue
        
        raise Exception("Could not decode text file with any supported encoding")
        
    except Exception as e:
        logger.error(f"Failed to extract text from TXT: {e}")
        raise Exception(f"Text file extraction failed: {e}")

def clean_text(text: str) -> str:
    """
    Clean and preprocess extracted text.
    
    Args:
        text: Raw extracted text
        
    Returns:
        Cleaned text content
    """
    try:
        # Basic text cleaning
        cleaned = text.strip()
        
        # Remove excessive whitespace
        cleaned = re.sub(r'\n\s*\n', '\n\n', cleaned)
        cleaned = re.sub(r' +', ' ', cleaned)
        
        # Remove control characters except newlines and tabs
        cleaned = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', cleaned)
        
        return cleaned
        
    except Exception as e:
        logger.warning(f"Text cleaning failed: {e}, returning original text")
        return text

def extract_text_from_file(file_content: bytes, filename: str) -> str:
    """
    Extract text from various file formats.
    
    Args:
        file_content: File content as bytes
        filename: Name of the file (used to determine format)
        
    Returns:
        Extracted and cleaned text content
        
    Raises:
        Exception: If extraction fails or file type is unsupported
    """
    try:
        file_extension = Path(filename).suffix.lower()
        
        logger.info(f"Extracting text from {filename} (extension: {file_extension})")
        
        # Route to appropriate extraction method
        if file_extension == '.pdf':
            raw_text = extract_text_from_pdf(file_content)
        elif file_extension == '.docx':
            raw_text = extract_text_from_docx(file_content)
        elif file_extension in ['.txt', '.md']:
            raw_text = extract_text_from_txt(file_content)
        else:
            raise Exception(f"Unsupported file type: {file_extension}")
        
        # Clean and preprocess the extracted text
        cleaned_text = clean_text(raw_text)
        
        logger.info(f"Text extraction completed: {len(cleaned_text)} characters")
        
        if len(cleaned_text.strip()) < 10:
            raise Exception("Extracted text is too short (minimum 10 characters required)")
        
        return cleaned_text
        
    except Exception as e:
        logger.error(f"Failed to extract text from {filename}: {e}")
        raise Exception(f"Text extraction failed for {filename}: {e}")

def get_supported_file_types() -> list:
    """
    Get list of supported file types for text extraction.
    
    Returns:
        List of supported file extensions
    """
    return ['.pdf', '.docx', '.txt', '.md']

def is_supported_file_type(filename: str) -> bool:
    """
    Check if the file type is supported for text extraction.
    
    Args:
        filename: Name of the file to check
        
    Returns:
        True if file type is supported, False otherwise
    """
    file_extension = Path(filename).suffix.lower()
    return file_extension in get_supported_file_types() 